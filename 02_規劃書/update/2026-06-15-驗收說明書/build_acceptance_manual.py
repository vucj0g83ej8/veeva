from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SCREENSHOT_DIR = ROOT / "screenshots"
OUTPUT = ROOT / "VeeVa會員管理系統_驗收說明書.docx"

FONT = "Microsoft JhengHei"
ACCENT = RGBColor(33, 107, 87)
DARK = RGBColor(18, 35, 32)
MUTED = RGBColor(96, 112, 108)
LIGHT_FILL = "EAF4EF"
TABLE_FILL = "F3F7F5"
BORDER = "D9E3DF"


FRONT_PAGES = [
    (
        "活動列表",
        "front-01-activities.png",
        "會員進入 LIFF 後的主要入口，可瀏覽全部活動、即將開始、已報名與已完成活動。",
        [
            "活動卡片顯示圖片、期間、標題、摘要與地點。",
            "已結束且使用者未參加的活動不會出現在主要列表。",
            "底部導覽可切換活動、最新資訊、兌換券與會員中心。",
        ],
    ),
    (
        "活動資訊頁",
        "front-02-activity-detail.png",
        "會員點擊活動後進入活動介紹頁，先閱讀活動內容，再執行報名、填問卷或分享。",
        [
            "不同活動型態會顯示不同主按鈕，例如我要報名或填寫問卷。",
            "已報名活動會顯示已報名並禁止重複報名。",
            "分享按鈕可透過 LINE 圖文卡片分享活動。",
        ],
    ),
    (
        "問卷活動介紹頁",
        "front-03-survey-activity.png",
        "問卷活動會先顯示活動資訊，會員點擊填寫問卷後才進入外部問卷頁。",
        [
            "問卷活動支援外部 OneTrust 問卷嵌入。",
            "完成任務後只會建立待確認紀錄，不會直接發放可使用兌換券。",
            "若活動有設定二段式獎勵，也會先進入待確認流程。",
        ],
    ),
    (
        "問卷填寫頁",
        "front-04-survey-page.png",
        "外部問卷嵌入在 LIFF 內顯示，會員可在系統內完成問卷任務。",
        [
            "問卷完成後會建立任務完成紀錄。",
            "已完成的問卷再次進入時會顯示已填寫完成。",
            "問卷 iframe 已針對手機寬度做縮放與水平溢出修正。",
        ],
    ),
    (
        "最新資訊列表",
        "front-05-news.png",
        "會員可瀏覽最新醫療資訊文章列表，並依文章分類快速閱讀。",
        [
            "文章卡片顯示圖片、日期、分類、標題與摘要。",
            "列表 UI 與活動頁統一，維持一致的閱讀體驗。",
            "點擊卡片可進入文章詳情頁。",
        ],
    ),
    (
        "文章詳情頁",
        "front-06-news-detail.png",
        "文章頁以完整閱讀版面呈現，支援後台自由排版內容。",
        [
            "文章頁不再固定顯示摘要或重點整理欄位。",
            "後台設定的粗體、字級、引用等格式會在前台呈現。",
            "會員可在文章底部點擊有幫助，系統會統計按讚數量。",
            "底部導覽保留，方便會員回到其他功能。",
        ],
    ),
    (
        "文章互動統計",
        "front-09-news-helpful.png",
        "會員閱讀文章後可點擊有幫助，前台會顯示目前累計的互動人數。",
        [
            "按讚數量由前台使用者互動累計，不由後台手動填寫。",
            "互動區固定在文章底部，會員閱讀完可直接操作。",
            "已登入會員點擊後可更新統計數量。",
        ],
    ),
    (
        "兌換券列表",
        "front-07-coupons.png",
        "會員可查看自己取得的兌換券，包含待確認、可兌換與已使用狀態。",
        [
            "待確認兌換券可被看見，但不可使用。",
            "人工確認後，狀態才會變成可兌換並取得兌換連結。",
            "可兌換券點擊後會進入兌換券詳情與確認使用流程。",
        ],
    ),
    (
        "會員中心",
        "front-08-member.png",
        "會員中心提供會員資料、邀請好友、邀請紀錄與客服協助入口。",
        [
            "可編輯姓名與 email。",
            "邀請好友會使用會員專屬分享連結。",
            "邀請紀錄會顯示透過會員連結加入的好友紀錄。",
        ],
    ),
]


ADMIN_PAGES = [
    (
        "儀表板",
        "admin-01-dashboard.png",
        "後台首頁提供問卷完成、待審核、審核通過與兌換券庫存概況。",
        [
            "可快速掌握待審核名單與整體狀態分布。",
            "左側導覽可進入權限、會員、活動、獎勵、最新資訊與兌換券管理。",
        ],
    ),
    (
        "權限管理",
        "admin-02-permissions.png",
        "管理後台管理者角色、啟用狀態與可操作權限。",
        [
            "可從 LINE 會員建立管理者。",
            "支援擁有者、管理員等角色與權限編輯。",
            "不顯示 LINE login ID，降低畫面資訊干擾。",
        ],
    ),
    (
        "會員管理",
        "admin-03-members.png",
        "查看已登入會員、會員狀態、角色標籤與會員設定。",
        [
            "會員列表顯示會員名稱、最後登入時間與操作。",
            "可編輯會員設定與發送兌換券。",
            "支援搜尋、分頁、待審核與已審核分頁管理。",
        ],
    ),
    (
        "活動管理",
        "admin-04-activities.png",
        "新增、編輯、預覽、啟停與封存活動。",
        [
            "支援問卷活動與活動報名類活動。",
            "可設定完成任務獎勵與邀請者二段式獎勵。",
            "活動圖片支援拖曳上傳與壓縮後存放 Firebase Storage。",
        ],
    ),
    (
        "獎勵發放",
        "admin-05-reward-distribution.png",
        "集中列出有獎勵規則的活動，人工確認後針對參加者發放獎勵。",
        [
            "完成任務後前台只會產生待確認券。",
            "管理者確認後才會轉成可兌換、分配序號與扣庫存。",
            "二段式邀請者獎勵會在人工發放時一併處理。",
        ],
    ),
    (
        "最新資訊管理",
        "admin-06-news.png",
        "管理文章清單、發布狀態、分類、來源與文章內容。",
        [
            "新增文章使用類 Word 的自由排版編輯器。",
            "支援文字粗體、斜體、字體大小與引用樣式。",
            "文章圖片可由編輯器直接上傳並插入內容。",
            "文章表格已調整為滿版寬度。",
        ],
    ),
    (
        "文章編輯器",
        "admin-09-news-editor.png",
        "新增或編輯文章時會開啟自由排版編輯器，可直接在後台完成文章內容排版。",
        [
            "編輯器提供粗體、斜體、字級、清單、引用、連結、圖片與分隔線工具。",
            "文章內容區支援長內容捲動，方便編輯較長文章。",
            "右側文章設定可管理列表摘要、來源、分類、發布日期、發布狀態、封面圖片與外部連結。",
        ],
    ),
    (
        "兌換券管理",
        "admin-07-rewards.png",
        "管理兌換券品項、分類、庫存、期限、圖片與序號連結。",
        [
            "可新增與編輯兌換券。",
            "支援拖曳上傳圖片與 Firebase Storage 預覽。",
            "支援匯入 Excel / CSV / TXT 兌換連結並逐一分配給會員。",
        ],
    ),
    (
        "系統設定",
        "admin-08-settings.png",
        "保留系統設定入口，用於後續管理全域設定與串接參數。",
        [
            "目前提供設定頁框架。",
            "後續如有額外新增功能，會在此頁新增。",
        ],
    ),
]


def set_east_asia_font(run, font=FONT):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def style_paragraph(paragraph, before=0, after=6, line=1.25):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_run(paragraph, text, size=11, bold=False, color=None):
    run = paragraph.add_run(text)
    set_east_asia_font(run)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    return run


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    style_paragraph(p, before=14 if level == 1 else 10, after=6)
    add_run(
        p,
        text,
        size=16 if level == 1 else 13,
        bold=True,
        color=ACCENT if level <= 2 else DARK,
    )
    return p


def add_body(doc, text, after=6):
    p = doc.add_paragraph()
    style_paragraph(p, after=after)
    add_run(p, text, size=11, color=DARK)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    style_paragraph(p, after=4)
    run = p.add_run(text)
    set_east_asia_font(run)
    run.font.size = Pt(10.5)
    run.font.color.rgb = DARK
    return p


def add_checklist_item(doc, text):
    p = doc.add_paragraph()
    style_paragraph(p, after=3)
    add_run(p, "□ ", size=11, bold=True, color=ACCENT)
    add_run(p, text, size=10.5, color=DARK)


def add_info_table(doc, rows, col_widths=None):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = col_widths or [Inches(2.0), Inches(4.3)]
    for row_index, row_data in enumerate(rows):
        row = table.rows[row_index]
        for col_index, value in enumerate(row_data):
            cell = row.cells[col_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            if row_index == 0:
                set_cell_fill(cell, TABLE_FILL)
            for paragraph in cell.paragraphs:
                style_paragraph(paragraph, after=0)
                run = paragraph.add_run(str(value))
                set_east_asia_font(run)
                run.font.size = Pt(9.5)
                run.font.bold = row_index == 0
                run.font.color.rgb = DARK
            if col_index < len(widths):
                cell.width = widths[col_index]
    doc.add_paragraph()
    return table


def add_screenshot_section(doc, title, image_name, description, bullets, image_width):
    add_heading(doc, title, level=2)
    add_body(doc, description, after=5)
    for item in bullets:
        add_bullet(doc, item)
    image_path = SCREENSHOT_DIR / image_name
    if image_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(p, before=4, after=3)
        run = p.add_run()
        run.add_picture(str(image_path), width=image_width)
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_paragraph(caption, after=10)
        add_run(caption, f"截圖：{title}", size=9, color=MUTED)
    else:
        add_body(doc, f"截圖檔案未找到：{image_name}", after=10)


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(10.5)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(header, "VeeVa 會員管理系統｜驗收說明書", size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(footer, "本文件供專案驗收與客戶確認使用", size=9, color=MUTED)
    return doc


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(p, before=36, after=4)
    add_run(p, "VeeVa 會員管理系統", size=24, bold=True, color=ACCENT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(p, after=18)
    add_run(p, "客戶驗收說明書", size=18, bold=True, color=DARK)

    add_info_table(
        doc,
        [
            ("項目", "內容"),
            ("前台 LIFF 網址", "https://vevva.web.app"),
            ("後台管理網址", "https://veeva-admin.web.app"),
            ("資料庫與儲存", "Firebase Firestore / Firebase Storage"),
            ("登入方式", "LINE Login / LIFF"),
            ("文件日期", date.today().strftime("%Y/%m/%d")),
        ],
        [Inches(1.7), Inches(4.8)],
    )

    add_body(
        doc,
        "本文件彙整目前已完成之 VeeVa 會員管理系統前台與後台功能，並附上正式站截圖、操作說明與驗收檢核項目，供客戶進行交付驗收。",
        after=12,
    )


def add_overview(doc):
    add_heading(doc, "一、系統驗收範圍", level=1)
    add_body(doc, "本次驗收範圍包含會員前台 LIFF App 與管理者後台平台。")
    for item in [
        "前台：活動列表、活動詳情、問卷填寫、最新資訊、文章詳情、兌換券、會員中心。",
        "後台：儀表板、權限管理、會員管理、活動管理、獎勵發放、最新資訊管理、兌換券管理、系統設定。",
        "資料：會員資料、活動紀錄、問卷完成紀錄、兌換券紀錄、邀請紀錄、文章與活動資料皆以 Firebase 為主要資料來源。",
        "獎勵：完成任務後前台顯示待確認兌換券，須由後台人工確認後才會轉為可使用。",
    ]:
        add_bullet(doc, item)

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "二、核心操作流程", level=1)
    flow_rows = [
        ("流程", "說明", "驗收重點"),
        ("LINE 登入", "會員從 LINE / LIFF 進入前台，系統檢查登入狀態。", "未登入時導向 LINE 登入；登入後記錄會員。"),
        ("活動參加", "會員進入活動詳情，依活動類型進行報名或填寫問卷。", "活動狀態、已報名與已完成顯示正確。"),
        ("問卷任務", "外部問卷嵌入 LIFF，會員可於系統內完成問卷。", "完成後建立任務完成紀錄，並進入待確認流程。"),
        ("待確認券", "完成任務後產生待確認兌換券。", "前台可見但不可使用，狀態顯示待確認。"),
        ("人工發放", "管理者於後台獎勵發放頁確認發放。", "發放後分配兌換連結、扣庫存並轉為可兌換。"),
        ("兌換使用", "會員開啟兌換券詳情，確認後使用優惠券。", "使用前跳出確認視窗，確認後狀態轉為已使用。"),
    ]
    add_info_table(doc, flow_rows, [Inches(1.25), Inches(3.0), Inches(2.1)])


def add_frontend(doc):
    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "三、前台 LIFF App 頁面說明", level=1)
    add_body(doc, "以下截圖取自正式站 https://vevva.web.app，以手機版 LIFF 使用情境呈現。")
    for item in FRONT_PAGES:
        add_screenshot_section(doc, *item, image_width=Inches(2.45))


def add_admin(doc):
    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "四、後台管理平台頁面說明", level=1)
    add_body(doc, "以下截圖以後台管理平台的桌面管理情境呈現，包含正式站與同版本驗收畫面。")
    for item in ADMIN_PAGES:
        add_screenshot_section(doc, *item, image_width=Inches(6.3))


def add_checklist(doc):
    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "五、驗收檢核表", level=1)
    groups = [
        (
            "前台會員功能",
            [
                "會員可透過 LINE Login 完成登入。",
                "活動列表、活動詳情與活動分類顯示正常。",
                "問卷活動可開啟外部問卷，手機版不會水平溢出。",
                "完成問卷後會記錄完成，並顯示已填寫完成。",
                "完成任務後產生待確認兌換券，且不可使用。",
                "最新資訊列表與文章詳情可正常瀏覽。",
                "會員中心可編輯姓名與 email。",
                "邀請好友與邀請紀錄功能可正常顯示。",
            ],
        ),
        (
            "後台管理功能",
            [
                "管理者可透過 LINE Login 進入後台。",
                "權限管理可新增與編輯管理者角色。",
                "會員管理可查看會員、編輯設定與發送兌換券。",
                "活動管理可新增、編輯、預覽、啟停與封存活動。",
                "獎勵發放頁可列出有獎勵規則的活動與參加者。",
                "管理者可人工確認發放待確認兌換券。",
                "最新資訊管理可使用自由排版編輯器編輯文章。",
                "兌換券管理可新增兌換券、上傳圖片與匯入兌換連結。",
            ],
        ),
        (
            "資料與部署",
            [
                "Firestore 會員、活動、兌換券與完成紀錄可正常讀寫。",
                "Firebase Storage 圖片上傳與前台顯示正常。",
                "前台正式網址 https://vevva.web.app 可正常開啟。",
                "後台正式網址 https://veeva-admin.web.app 可正常開啟。",
                "GitHub 已備份最新版本。",
            ],
        ),
    ]
    for title, checks in groups:
        add_heading(doc, title, level=2)
        for check in checks:
            add_checklist_item(doc, check)

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "六、驗收簽核", level=1)
    add_info_table(
        doc,
        [
            ("角色", "姓名 / 簽名", "日期", "備註"),
            ("客戶驗收", "", "", ""),
            ("專案交付", "", "", ""),
        ],
        [Inches(1.4), Inches(2.2), Inches(1.3), Inches(1.6)],
    )


def main():
    doc = setup_document()
    add_cover(doc)
    add_overview(doc)
    add_frontend(doc)
    add_admin(doc)
    add_checklist(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
