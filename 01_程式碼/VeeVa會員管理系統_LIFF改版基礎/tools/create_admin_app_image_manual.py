from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
OUT_DIR.mkdir(exist_ok=True)
OUTPUT = OUT_DIR / "VeeVa_Admin_後台管理_App_圖片說明書.docx"

GREEN = RGBColor(33, 107, 87)
DARK = RGBColor(24, 36, 31)
MUTED = RGBColor(93, 107, 101)


PAGES = [
    (
        "儀表板",
        "/Users/ouyangtaisen/Desktop/截圖 2026-05-08 上午11.06.06.png",
        "後台首頁總覽問卷完成數、待審核數、審核通過數、兌換券庫存與名單狀態分布。",
    ),
    (
        "會員管理",
        "/Users/ouyangtaisen/Desktop/截圖 2026-05-08 上午11.06.14.png",
        "整合待審核與已審核會員名單，管理員可切換分頁並執行審核通過操作。",
    ),
    (
        "活動管理",
        "/Users/ouyangtaisen/Desktop/截圖 2026-05-08 上午11.06.21.png",
        "管理活動名稱、狀態、活動期間與備註，並提供新增活動入口。",
    ),
    (
        "兌換券管理",
        "/Users/ouyangtaisen/Desktop/截圖 2026-05-08 上午11.06.32.png",
        "管理兌換券庫存、發放數、兌換數、期限、狀態與補庫存操作。",
    ),
    (
        "最新資訊管理",
        "/Users/ouyangtaisen/Desktop/截圖 2026-05-08 上午11.06.40.png",
        "管理醫療資訊的標題、狀態、發布日期與分類，並提供新增資訊入口。",
    ),
]


def set_run_font(run, size=None, bold=None, color=None):
    font_name = "Microsoft JhengHei"
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def setup_doc(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    for style_name in ("Normal", "Title", "Heading 1"):
        style = doc.styles[style_name]
        style.font.name = "Microsoft JhengHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")

    doc.styles["Normal"].font.size = Pt(11)
    doc.styles["Title"].font.size = Pt(28)
    doc.styles["Title"].font.bold = True


def add_cover(doc):
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("VeeVa Admin")
    set_run_font(run, size=30, bold=True, color=DARK)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("後台管理 App 圖片說明書")
    set_run_font(run, size=22, bold=True, color=GREEN)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Prototype 版｜2026/05/08")
    set_run_font(run, size=12, color=MUTED)

    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro.paragraph_format.space_before = Pt(20)
    run = intro.add_run("本文件直接套用後台管理 App 截圖，作為後台功能展示、需求確認與開發驗收參考。")
    set_run_font(run, size=12, color=DARK)
    doc.add_page_break()


def add_image_page(doc, index, title, image_path, note):
    heading = doc.add_paragraph()
    heading.paragraph_format.space_after = Pt(0)
    run = heading.add_run(f"{index}. {title}")
    set_run_font(run, size=18, bold=True, color=GREEN)

    desc = doc.add_paragraph()
    desc.paragraph_format.space_after = Pt(6)
    run = desc.add_run(note)
    set_run_font(run, size=10.5, color=MUTED)

    image_para = doc.add_paragraph()
    image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_para.paragraph_format.space_after = Pt(2)
    image_para.add_run().add_picture(image_path, width=Inches(9.9))

    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(f"圖 {index}　{title}")
    set_run_font(run, size=9.5, color=MUTED)

    if index != len(PAGES):
        doc.add_page_break()


def main():
    doc = Document()
    setup_doc(doc)
    add_cover(doc)
    for index, (title, image_path, note) in enumerate(PAGES, 1):
        add_image_page(doc, index, title, image_path, note)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
