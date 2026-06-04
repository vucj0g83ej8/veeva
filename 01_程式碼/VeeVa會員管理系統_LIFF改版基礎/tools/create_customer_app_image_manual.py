from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
OUT_DIR.mkdir(exist_ok=True)
OUTPUT = OUT_DIR / "VeeVa_顧客端_App_圖片說明書.docx"

GREEN = RGBColor(33, 107, 87)
DARK = RGBColor(24, 36, 31)
MUTED = RGBColor(93, 107, 101)


PAGES = [
    ("活動消息", "/Users/ouyangtaisen/Desktop/截圖 2026-05-08 上午10.15.22.png", "查看目前活動，從「立即開始」進入問卷流程。"),
    ("會員登入", "/Users/ouyangtaisen/Desktop/截圖 2026-05-08 上午10.15.33.png", "提供 LINE 與 Google 登入。"),
    ("Veeva 問卷填寫", "/Users/ouyangtaisen/Desktop/截圖 2026-05-08 上午10.32.48.png", "在 App 內嵌 Veeva Consent Form，完成後送出審核。"),
    ("送出成功", "/Users/ouyangtaisen/Desktop/截圖 2026-05-08 上午10.35.07.png", "顯示 Thank You 與待審核狀態。"),
    ("會員中心", "/Users/ouyangtaisen/Desktop/截圖 2026-05-08 上午10.35.28.png", "顯示會員資料、資格狀態與會員功能。"),
    ("系統訊息", "/Users/ouyangtaisen/Desktop/截圖 2026-05-08 上午10.35.46.png", "點擊右上角鈴鐺查看系統通知。"),
    ("最新資訊", "/Users/ouyangtaisen/Desktop/截圖 2026-05-08 上午10.35.59.png", "條列呈現醫療與活動相關資訊。"),
    ("兌換券", "/Users/ouyangtaisen/Desktop/截圖 2026-05-08 上午10.36.08.png", "查看可用兌換券與兌換期限。"),
]


def set_run_font(run, size=None, bold=None, color=None):
    font_name = "Microsoft JhengHei"
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def setup_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    for style_name in ("Normal", "Title", "Heading 1"):
        style = doc.styles[style_name]
        style.font.name = "Microsoft JhengHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")

    doc.styles["Normal"].font.size = Pt(11)
    doc.styles["Title"].font.size = Pt(28)
    doc.styles["Title"].font.bold = True


def add_cover(doc):
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("VeeVa 顧客端 App")
    set_run_font(run, size=28, bold=True, color=DARK)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("圖片操作說明書")
    set_run_font(run, size=20, bold=True, color=GREEN)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Prototype 版｜2026/05/08")
    set_run_font(run, size=12, color=MUTED)

    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro.paragraph_format.space_before = Pt(20)
    run = intro.add_run("本文件直接套用顧客端 App 截圖，作為頁面流程與功能展示用說明書。")
    set_run_font(run, size=11.5, color=DARK)
    doc.add_page_break()


def add_image_page(doc, index, title, image_path, note):
    heading = doc.add_paragraph()
    heading.paragraph_format.space_after = Pt(2)
    run = heading.add_run(f"{index}. {title}")
    set_run_font(run, size=18, bold=True, color=GREEN)

    desc = doc.add_paragraph()
    desc.paragraph_format.space_after = Pt(10)
    run = desc.add_run(note)
    set_run_font(run, size=10.5, color=MUTED)

    image_para = doc.add_paragraph()
    image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_para.paragraph_format.space_after = Pt(4)
    image_para.add_run().add_picture(image_path, width=Inches(3.25))

    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption.add_run(f"圖 {index}　{title}")
    set_run_font(run, size=9.5, color=MUTED)

    if index != len(PAGES):
        doc.add_page_break()


def main():
    doc = Document()
    setup_doc(doc)
    add_cover(doc)
    for index, (title, image_path, note) in enumerate(PAGES, 1):
        add_image_page(doc, index, title, image_path, note)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
