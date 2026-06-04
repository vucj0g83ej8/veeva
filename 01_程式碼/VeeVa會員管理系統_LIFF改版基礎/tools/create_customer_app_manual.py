from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
OUT_DIR.mkdir(exist_ok=True)
OUTPUT = OUT_DIR / "VeeVa_顧客端_App_操作說明書.docx"

GREEN = RGBColor(33, 107, 87)
DARK = RGBColor(24, 36, 31)
MUTED = RGBColor(93, 107, 101)
LIGHT_GREEN = "EAF3EA"
LIGHT_GRAY = "F5F7F8"
BORDER = "DDE5DF"


SCREENSHOTS = [
    {
        "title": "活動消息",
        "path": "/Users/ouyangtaisen/Desktop/截圖 2026-05-08 上午10.15.22.png",
        "purpose": "使用者進入 App 後可瀏覽目前可參加的活動，例如問卷換券、研討會提醒與院所任務。",
        "steps": [
            "從底部導覽列點選「活動」。",
            "查看各活動卡片的狀態、獎勵與開始條件。",
            "點選「立即開始」進入登入與問卷流程。",
        ],
        "notes": "活動卡片可持續新增，適合放置短期任務、院所限定任務與未來活動預告。",
    },
    {
        "title": "會員登入",
        "path": "/Users/ouyangtaisen/Desktop/截圖 2026-05-08 上午10.15.33.png",
        "purpose": "會員登入頁提供 LINE 與 Google 兩種登入方式，畫面保持簡潔並支援返回上一頁。",
        "steps": [
            "從活動頁點選活動後進入登入畫面。",
            "選擇「使用 LINE 登入」或「使用 Google 登入」。",
            "登入成功後進入 Veeva 問卷填寫頁。",
        ],
        "notes": "正式版可接入 LINE Login 與 Google Sign-In，並取得會員基本資料。",
    },
    {
        "title": "Veeva 問卷填寫",
        "path": "/Users/ouyangtaisen/Desktop/截圖 2026-05-08 上午10.32.48.png",
        "purpose": "問卷頁直接嵌入 Veeva Consent Form，使用者可在 App 內完成外部表單。",
        "steps": [
            "閱讀 Veeva Consent Form 內容。",
            "依表單要求完成填寫或同意流程。",
            "填寫完成後點選「我已完成問卷，送出審核」。",
        ],
        "notes": "目前以 iframe 嵌入 OneTrust / Veeva 表單，正式版可視供應商支援情況加入 callback 或 webhook。",
    },
    {
        "title": "問卷送出成功",
        "path": "/Users/ouyangtaisen/Desktop/截圖 2026-05-08 上午10.35.07.png",
        "purpose": "送出後顯示 Thank You 卡片，提醒使用者目前狀態為待審核。",
        "steps": [
            "確認畫面顯示 Thank You。",
            "查看目前狀態是否為「待審核」。",
            "點選「回首頁」返回後續流程。",
        ],
        "notes": "目前按鈕功能暫時沿用模擬審核通過流程，正式版可改為返回活動首頁或會員中心。",
    },
    {
        "title": "會員中心",
        "path": "/Users/ouyangtaisen/Desktop/截圖 2026-05-08 上午10.35.28.png",
        "purpose": "會員中心顯示會員基本資料、資格狀態、已得券與已邀請數，並提供常用會員功能。",
        "steps": [
            "從底部導覽列點選「會員」。",
            "確認會員姓名、院所、科別與驗證狀態。",
            "使用會員功能，例如編輯資料、通知設定、活動紀錄與客服協助。",
        ],
        "notes": "會員功能可作為後續 CRM、通知偏好與活動追蹤的入口。",
    },
    {
        "title": "系統訊息",
        "path": "/Users/ouyangtaisen/Desktop/截圖 2026-05-08 上午10.35.46.png",
        "purpose": "使用者可透過右上角通知按鈕查看系統發送的訊息。",
        "steps": [
            "點選右上角鈴鐺圖示。",
            "查看系統訊息彈窗。",
            "確認問卷、咖啡券發送、會員資料更新與活動上架提醒。",
        ],
        "notes": "訊息可串接通知中心，未來支援已讀狀態、推播與訊息分類。",
    },
    {
        "title": "最新資訊",
        "path": "/Users/ouyangtaisen/Desktop/截圖 2026-05-08 上午10.35.59.png",
        "purpose": "最新資訊頁以條列卡片呈現醫療與產品相關消息，方便會員快速瀏覽。",
        "steps": [
            "從底部導覽列點選「最新資訊」。",
            "瀏覽消息標題、日期與來源。",
            "點選卡片可進入消息詳情或外部來源。",
        ],
        "notes": "後台可建立最新資訊管理，支援發布、草稿、分類與排序。",
    },
    {
        "title": "兌換券",
        "path": "/Users/ouyangtaisen/Desktop/截圖 2026-05-08 上午10.36.08.png",
        "purpose": "兌換券頁列出可用券與兌換期限，使用者可直接進行兌換。",
        "steps": [
            "從底部導覽列點選「兌換券」。",
            "查看商品名稱與兌換期限。",
            "點選「兌換」後確認兌換內容。",
        ],
        "notes": "兌換券可包含咖啡、飲品、購物金、實體贈品與會員點數。",
    },
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "nil")


def set_run_font(run, size=None, bold=None, color=None):
    font_name = "Microsoft JhengHei"
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_paragraph(doc, text="", style=None, size=11, bold=False, color=None, after=6):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_bullet(cell, text):
    p = cell.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=10.5)
    return p


def add_section_title(doc, title, subtitle=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(title)
    set_run_font(run, size=16, bold=True, color=GREEN)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(12)
        run2 = p2.add_run(subtitle)
        set_run_font(run2, size=10.5, color=MUTED)


def setup_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    for style_name in ("Normal", "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3", "List Bullet"):
        style = styles[style_name]
        style.font.name = "Microsoft JhengHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")

    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.25
    styles["Title"].font.size = Pt(26)
    styles["Title"].font.bold = True
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = GREEN
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.bold = True
    styles["Heading 2"].font.color.rgb = GREEN


def add_header_footer(doc):
    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.text = ""
        left = header.add_run("VeeVa 顧客端 App 操作說明書")
        set_run_font(left, size=9.5, color=MUTED)
        header.alignment = WD_ALIGN_PARAGRAPH.LEFT
        footer = section.footer.paragraphs[0]
        footer.text = ""
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("VeeVa Member App Guide")
        set_run_font(run, size=9, color=MUTED)


def add_cover(doc):
    add_paragraph(doc, "VeeVa Member App", size=12, bold=True, color=GREEN, after=16)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("顧客端 App 操作說明書")
    set_run_font(run, size=28, bold=True, color=DARK)
    add_paragraph(doc, "適用對象：使用者、客戶成功團隊、專案驗收與教育訓練", size=12, color=MUTED, after=28)

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(table)
    labels = [("版本", "Prototype"), ("日期", "2026/05/08"), ("內容", "顧客端主要流程")]
    for cell, (label, value) in zip(table.rows[0].cells, labels):
        set_cell_shading(cell, LIGHT_GREEN)
        set_cell_border(cell, "CFE7DC")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(label)
        set_run_font(r1, size=9.5, bold=True, color=GREEN)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(value)
        set_run_font(r2, size=11, bold=True, color=DARK)

    add_paragraph(doc, "", after=18)
    callout = doc.add_table(rows=1, cols=1)
    set_cell_shading(callout.cell(0, 0), LIGHT_GRAY)
    set_cell_border(callout.cell(0, 0), BORDER)
    cell = callout.cell(0, 0)
    p = cell.paragraphs[0]
    r = p.add_run("文件用途")
    set_run_font(r, size=12, bold=True, color=GREEN)
    p2 = cell.add_paragraph()
    r2 = p2.add_run("本說明書整理顧客端 App 的主要操作畫面與使用流程，作為內部溝通、客戶展示、教育訓練與後續開發驗收參考。")
    set_run_font(r2, size=11, color=DARK)
    doc.add_page_break()


def add_overview(doc):
    add_section_title(doc, "一、操作流程總覽", "顧客端 App 以活動參與、問卷填寫、資格審核、兌換券與會員中心為主流程。")
    steps = [
        ("1", "活動消息", "瀏覽可參加活動並啟動任務"),
        ("2", "會員登入", "使用 LINE 或 Google 登入會員"),
        ("3", "Veeva 問卷", "在 App 內嵌頁面完成同意與問卷"),
        ("4", "送出審核", "送出後進入待審核狀態"),
        ("5", "會員中心", "查看會員資料、資格狀態與功能入口"),
        ("6", "兌換券", "查看可用券並確認兌換"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header = table.rows[0].cells
    for idx, text in enumerate(["順序", "流程", "說明"]):
        set_cell_shading(header[idx], LIGHT_GREEN)
        set_cell_border(header[idx], "CFE7DC")
        r = header[idx].paragraphs[0].add_run(text)
        set_run_font(r, size=10.5, bold=True, color=GREEN)
    for num, flow, desc in steps:
        row = table.add_row().cells
        for idx, text in enumerate([num, flow, desc]):
            set_cell_border(row[idx])
            if idx == 0:
                row[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = row[idx].paragraphs[0].add_run(text)
            set_run_font(r, size=10.5, bold=(idx == 1), color=DARK)
    add_paragraph(doc, "底部導覽列主要包含「活動」「最新資訊」「兌換券」「會員」，右上角鈴鐺提供系統訊息。", size=10.5, color=MUTED, after=0)
    doc.add_page_break()


def add_feature_page(doc, index, item):
    add_section_title(doc, f"{index}. {item['title']}", item["purpose"])
    layout = doc.add_table(rows=1, cols=2)
    layout.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(layout)
    text_cell, image_cell = layout.rows[0].cells
    text_cell.width = Inches(3.55)
    image_cell.width = Inches(2.85)
    text_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    image_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    p = text_cell.paragraphs[0]
    r = p.add_run("操作步驟")
    set_run_font(r, size=12, bold=True, color=GREEN)
    for step in item["steps"]:
        add_bullet(text_cell, step)

    p2 = text_cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(10)
    r2 = p2.add_run("設計與開發備註")
    set_run_font(r2, size=12, bold=True, color=GREEN)
    p3 = text_cell.add_paragraph()
    p3.paragraph_format.line_spacing = 1.25
    r3 = p3.add_run(item["notes"])
    set_run_font(r3, size=10.5, color=DARK)

    image_para = image_cell.paragraphs[0]
    image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_para.add_run().add_picture(item["path"], width=Inches(2.65))
    caption = image_cell.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = caption.add_run(f"圖 {index}　{item['title']}")
    set_run_font(rc, size=9.5, color=MUTED)

    if index != len(SCREENSHOTS):
        doc.add_page_break()


def build_doc():
    doc = Document()
    setup_styles(doc)
    add_header_footer(doc)
    add_cover(doc)
    add_overview(doc)
    for idx, item in enumerate(SCREENSHOTS, 1):
        add_feature_page(doc, idx, item)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_doc()
